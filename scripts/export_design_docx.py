from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Inches, Pt

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))


def _set_base_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def export_docx(md_path, docx_path):
    doc = Document()
    _set_base_style(doc)

    in_code = False
    code_lines = []
    image_re = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")

    for raw in Path(md_path).read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                para = doc.add_paragraph("\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        image_match = image_re.match(line.strip())
        if image_match:
            caption, rel_path = image_match.groups()
            img_path = Path(rel_path)
            if not img_path.is_absolute():
                img_path = ROOT_DIR / img_path
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(5.8))
                if caption:
                    doc.add_paragraph(caption)
            else:
                doc.add_paragraph(f"[Missing image] {rel_path}")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line.strip() == "":
            doc.add_paragraph("")
            continue
        doc.add_paragraph(line)

    doc.save(docx_path)


def main():
    md_path = ROOT_DIR / "DESIGN.md"
    out_path = ROOT_DIR / "Design Document.docx"
    export_docx(md_path, out_path)
    print(f"[done] {out_path}")


if __name__ == "__main__":
    main()
